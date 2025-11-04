import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import math

def row_has_data(row, cols):
    for col in cols:
        val = row[col]
        if pd.notnull(val):
            v = str(val).strip()
            if v != '' and v.lower() != 'nan':
                # If value is numeric or looks like a percentage string (e.g., '98%')
                if v.endswith('%'):
                    return True
                try:
                    float(v.replace('%',''))
                    return True
                except:
                    pass
    return False


# Helper functions for parsing and formatting values
def parse_cell_value(value):
    if isinstance(value, str):
        v = value.strip()
        if v.endswith('%'):
            try:
                return float(v.replace('%', '')) / 100, 'percent'
            except ValueError:
                return None, 'other'
        if '%' in v:  # e.g., "98.5 %"
            nums = ''.join(c for c in v if (c.isdigit() or c == '.' or c == '-'))
            try:
                return float(nums) / 100, 'percent'
            except ValueError:
                return None, 'other'
    try:
        num = float(value)
        if 0 <= num <= 1:
            return num, 'maybe_percent'
        return num, 'numeric'
    except (ValueError, TypeError):
        return None, 'other'

def format_stat(value, stat_type):
    if value is None:
        return 'N/A'
    if stat_type == 'percent' or stat_type == 'maybe_percent':
        return f"{round(value * 100)}%"
    else:
        return f"{round(value, 2)}"

# Read Excel file (update path if needed)
file_path = 'MTC_MUMBAI.xlsx'
df = pd.read_excel(file_path, sheet_name='MTC_Mumbai', header=None)

section_titles = [
    'Client Service Delivery Indicator',
    'Therapist Performance Indicator',
    'Frondesk Administration/ Business Indicators',
    'Offline Marketing Plan/ Events',
    'Outcomes'
]

# Identify section start indices
section_indices = {}
for idx, row in df.iterrows():
    if pd.notnull(row[2]):
        val = str(row[2]).strip()
        if val in section_titles:
            section_indices[val] = idx

sorted_sections = sorted(section_indices.items(), key=lambda x: x[1])
section_slices = {}
for i, (section, start_idx) in enumerate(sorted_sections):
    end_idx = sorted_sections[i+1][1] if i+1 < len(sorted_sections) else len(df)
    section_slices[section] = (start_idx + 1, end_idx)

months_cols = [3, 4, 5, 6, 7]
month_names = ['April', 'May', 'June', 'July', 'August']

# Prompt user for single or all month-to-month stats
print("Select month-to-month comparison option:")
print("1: All months (April-May, May-June, June-July, July-August)")
print("2: Single month comparison")
choice = input("Enter choice (1 or 2): ").strip()

single_month_index = None
if choice == '2':
    print("Enter the month pair you want stats for (e.g., 'June to July'):")
    user_input = input().strip().lower()
    month_pairs = {
        'april to may': 0,
        'may to june': 1,
        'june to july': 2,
        'july to august': 3,
    }
    if user_input in month_pairs:
        single_month_index = month_pairs[user_input]
    else:
        print("Invalid month pair input. Defaulting to all months.")
        choice = '1'

# Create document and apply formatting
doc = Document()
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(12)
style.paragraph_format.line_spacing = 1.15

h1_style = doc.styles['Heading 1']
h1_font = h1_style.font
h1_font.name = 'Calibri'
h1_font.size = Pt(14)
h1_font.bold = True
h1_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

for heading in ['Heading 2', 'Heading 3']:
    h_style = doc.styles[heading]
    h_font = h_style.font
    h_font.name = 'Calibri'
    h_font.size = Pt(12)
    h_font.bold = True
    h_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

title = doc.add_heading('Month-to-Month Comparison Report', level=0)
title.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

section = doc.sections[0]
footer = section.footer
paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
run = paragraph.add_run()
fldChar1 = OxmlElement('w:fldChar')
fldChar1.set(qn('w:fldCharType'), 'begin')
instrText = OxmlElement('w:instrText')
instrText.text = 'PAGE'
fldChar2 = OxmlElement('w:fldChar')
fldChar2.set(qn('w:fldCharType'), 'end')
run._r.append(fldChar1)
run._r.append(instrText)
run._r.append(fldChar2)

for section, (start, end) in section_slices.items():
    doc.add_heading(section, level=1).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    section_df = df.iloc[start:end]
    for _, row in section_df.iterrows():
        indicator = row[2]
        if pd.isna(indicator) or str(indicator).strip() == '' or str(indicator).strip() in section_titles:
            continue

        if not row_has_data(row, months_cols):
    # Add heading for indicator
            doc.add_heading(f"Statistics for '{indicator}':", level=3).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    # Add info paragraph
            p = doc.add_paragraph()
            p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            p.add_run("No values present in the period April to August to show statistics.").italic = True
    # Skip to next row
            continue


        # Parse and detect value types
        row_values = []
        types = []
        for val in row[months_cols]:
            parsed_val, val_type = parse_cell_value(val)
            row_values.append(parsed_val)
            types.append(val_type)

        if not any(v is not None for v in row_values):
            valid_values = [v for v in row_values if v is not None]

            if len(valid_values) == 0:
                doc.add_heading(f"Statistics for '{indicator}':", level=3).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p.add_run("No values present in the available months to show statistics.").italic = True
                continue
            # else continue processing on valid_values below


        # Stat type detection: if all are percent/maybe_percent, treat as percent
        if all(t in ('percent', 'maybe_percent') for t in types if t != 'other'):
            display_type = 'percent'
        else:
            display_type = 'numeric'

        doc.add_heading(f"Statistics for '{indicator}':", level=3).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # Select month pairs
        month_indices = [single_month_index] if single_month_index is not None else range(len(months_cols) - 1)

        for i in month_indices:
            prev_val = row_values[i]
            curr_val = row_values[i + 1]

            if prev_val is not None and curr_val is not None:
                difference = curr_val - prev_val
                pct_change = ((difference) / prev_val) * 100 if prev_val != 0 else None
                moving_avg = (prev_val + curr_val) / 2
                ratio = (curr_val / prev_val) if prev_val != 0 else None
                median = sorted([prev_val, curr_val])[1]
                std_dev = pd.Series([prev_val, curr_val]).std()

                bullet = doc.add_paragraph(f"{month_names[i+1]} vs {month_names[i]}:", style='List Bullet')
                bullet.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                p = doc.add_paragraph()
                p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                p.add_run(f"{month_names[i]} Value: ").bold = True
                p.add_run(f"{format_stat(prev_val, display_type)}\n")
                p.add_run(f"{month_names[i+1]} Value: ").bold = True
                p.add_run(f"{format_stat(curr_val, display_type)}\n")
                p.add_run(f"Difference: ").bold = True
                p.add_run(f"{format_stat(difference, display_type)}\n")
                p.add_run(f"Percentage Change: ").bold = True
                if pct_change is not None and not math.isnan(pct_change):
                    p.add_run(f"{round(pct_change)}%\n")
                else:
                    p.add_run("N/A (division by zero or missing data)\n")
                p.add_run(f"Moving Average: ").bold = True
                p.add_run(f"{format_stat(moving_avg, display_type)}\n")
                if ratio is not None:
                    p.add_run(f"Ratio (Current/Previous): ").bold = True
                    p.add_run(f"{format_stat(ratio, display_type)}\n")
                else:
                    p.add_run(f"Ratio (Current/Previous): ").bold = True
                    p.add_run("N/A (division by zero)\n")
                p.add_run(f"Median: ").bold = True
                p.add_run(f"{format_stat(median, display_type)}\n")
                p.add_run(f"Standard Deviation: ").bold = True
                p.add_run(f"{format_stat(std_dev, display_type)}\n")

                if pct_change is not None:
                    if pct_change > 10:
                        insight = "Significant growth observed."
                    elif 0 < pct_change <= 10:
                        insight = "Moderate increase noted."
                    elif -10 <= pct_change <= 0:
                        insight = "Slight decline observed; monitor closely."
                    else:
                        insight = "Considerable decrease; requires immediate attention."
                    insight_para = doc.add_paragraph(f"Insight: {insight}", style='Intense Quote')
                    insight_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                else:
                    insight_para = doc.add_paragraph("Insight: Percentage change not available due to zero baseline.", style='Intense Quote')
                    insight_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

                doc.add_paragraph()  # spacing
            else:
                missing_para = doc.add_paragraph(f"{month_names[i+1]} vs {month_names[i]}: Data missing for comparison", style='List Bullet')
                missing_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                doc.add_paragraph()
                # After the 'for i in month_indices:' loop ends (still inside the per-row processing)
                missing_months = [month_names[idx] for idx, val in enumerate(row_values) if val is None]

                if missing_months:
                    note_text = f"Note: The statistics above do not include data for the following month(s): {', '.join(missing_months)}."
                    note_para = doc.add_paragraph(note_text)
                    note_para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    note_para.runs[0].italic = True
                    doc.add_paragraph()  # Add extra blank line for spacing


# Save the report
output_filename = 'Professional1_July_to_August_Comparison_Report.docx'
doc.save(output_filename)
print(f"Report saved as {output_filename}")
